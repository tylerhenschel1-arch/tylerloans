"""
Daily blog draft generator — writes to _drafts/ for Tyler review.

Pipeline:
  1. Pick today's topic from the rotating list
  2. Generate the blog post body via Claude (with full CMG compliance rules)
  3. Generate 3 FAQ items
  4. Generate a 60-second video script
  5. Write _drafts/<YYYY-MM-DD>_<slug>/
       - index.html  -> pre-rendered blog page (publisher will move it to blog/)
       - review.docx -> Tyler reviews this in Word/Pages; contains video script + body
  6. Commit + push when running in CI

Tyler's workflow after this script runs:
  1. Pull the latest _drafts/
  2. Open review.docx -> proofread, run through CMG AI
  3. Record video reading the 60-second script
  4. Upload to Submagic -> get final video -> upload to YouTube
  5. Run: bin/publish <slug> --video <youtube-url>

The publisher script handles moving the pre-rendered HTML into blog/,
inserting the YouTube embed, updating posts.json + sitemap.xml,
and committing the live version.
"""

import anthropic
import os
import json
import datetime
import re
import subprocess

client = anthropic.Anthropic(api_key=os.environ['ANTHROPIC_API_KEY'])

TOPICS = [
    # First-time buyers — evergreen high volume
    "First time homebuyer mistakes to avoid in Houston TX in 2026",
    "How much house can you really afford in the Houston area",
    "FHA vs conventional loans in Houston — which is better for you",
    "VA loans explained for veterans and active duty buyers in Texas",
    "USDA loans in Texas — the hidden zero-down mortgage option",
    "What credit score do you need to buy a home in Texas",
    "How mortgage rates actually work — a Houston buyer's guide",
    "Why two buyers with the same credit score get different rates",
    "The complete mortgage pre-approval checklist for Houston homebuyers",
    "What underwriters really look for — Houston mortgage guide",
    "First time home buyer programs in the Houston metro 2026",
    "Down payment assistance programs in Texas — complete guide",
    "TDHCA loan programs explained for Texas first time buyers",
    "Grants vs down payment assistance — what is the difference in Texas",
    "Can parents help with your down payment in Texas — gift fund rules explained",
    "How much money do you need to buy a house in Houston",
    "How to prepare financially one year before buying a home in Texas",
    "The real cost of waiting to buy a home in Houston",
    "Renting vs buying in Houston — the real numbers in 2026",
    "Why waiting for rates to drop might backfire for Houston buyers",
    "Is 2026 a good time to buy a home in the Houston area",
    "How to compete against cash buyers in Houston TX",
    "What closing costs should buyers expect in Texas",
    "Hidden costs of homeownership nobody talks about in Texas",
    "Can you buy a house with bad credit in Texas",
    "Can gig workers qualify for mortgages in Houston",
    "Self-employed mortgage guide for Houston business owners",
    "How bonuses commission and overtime income are calculated for mortgages",
    "Can you buy a home with student loans in Texas",
    "Co-signing a mortgage in Texas — risks and benefits explained",
    "Should you pay off debt before buying a home in Houston",
    "How to buy a home with low savings in Texas",
    "Buying a home after bankruptcy in Texas — what you need to know",
    "Buying a home after divorce in Texas — a complete guide",
    "Mortgage myths that cost Houston buyers money",
    "The psychology of buying your first house",
    "Top questions buyers should ask their mortgage lender in Houston",
    "What I wish every Houston homebuyer knew before starting the process",

    # Houston metro neighborhoods — distributed across the MSA
    "Buying a home in Tomball TX — neighborhood guide and mortgage tips",
    "Buying a home in Cypress TX — what buyers need to know",
    "Buying a home in Sugar Land TX — neighborhood and financing guide",
    "Buying a home in The Woodlands TX — mortgage and market guide",
    "Buying a home in Pearland TX — what buyers should expect",
    "Buying a home in Spring TX — neighborhood and mortgage primer",
    "Buying a home in Magnolia TX — affordability and financing guide",
    "Buying a home in Conroe TX — mortgage and lifestyle guide",
    "Buying a home in Friendswood TX — a Houston buyer's perspective",
    "Buying a home in Memorial Houston — luxury market financing",
    "Buying a home in the Heights Houston — financing and considerations",
    "Buying a home in Bellaire Houston — neighborhood and mortgage guide",
    "Buying a home in Katy TX — neighborhood and mortgage guide",

    # Texas taxes and escrow — SEO goldmine
    "Understanding MUD taxes in Texas — what every Houston buyer should know",
    "Why are Texas property taxes so high and what buyers can do",
    "How property tax protests work in Harris County and Fort Bend",
    "When does your Texas escrow account go up and why",
    "What is a tax certificate and why your Houston builder should provide one",
    "How appraisal districts assess your home in Texas",
    "Why your Houston mortgage payment changed this year",
    "Texas homestead exemption — what every homeowner needs to know",
    "How to read your Texas property tax bill correctly",
    "Why MUD districts exist in Katy and Cypress TX",
    "What is a PID and how does it affect your Houston home",
    "Why escrow shortages happen and how to avoid them",
    "When can you remove escrow from your mortgage in Texas",

    # Investors and rental property
    "DSCR loans for Houston investors — the complete 2026 guide",
    "How to buy your first rental property in the Houston area",
    "Should you use a HELOC or refinance to buy investment property",
    "How investors use the AIO loan to scale faster in Texas",
    "What is a 1031 exchange and how Houston investors use it",
    "Investment property loan requirements in Texas 2026",
    "How to structure your first 5 rental properties in Houston",
    "Short term rental loans for Galveston and Houston investors",
    "How DSCR loans calculate rental income in Texas",
    "Conventional vs DSCR investment property loans — which to choose",
    "Why Texas is a hot market for out of state real estate investors",
    "How to refinance your Houston rental property in 2026",
    "What is portfolio lending and how Houston investors benefit",
    "How appraisals work on investment properties in Texas",

    # All-In-One loan — Tyler's signature product
    "What is the AIO loan and how does it work for Houston homeowners",
    "All in One mortgage vs HELOC — which is right for your situation",
    "How business owners in Houston can save big with the AIO loan",
    "Why high earners benefit most from the All In One mortgage",
    "How the AIO loan reduces interest using your daily cash flow",
    "Self employed Houston homeowners — why AIO might be your best fit",
    "Is the All In One mortgage right for Houston physicians and dentists",
    "How real estate investors in Houston use the AIO for liquidity (on a primary or second home)",
    "AIO loan calculator — how much can you save in Houston",
    "How to qualify for an AIO loan in Texas",
    "What kinds of income work best for AIO loans",
    "AIO loan rates vs HELOC rates — what to know in 2026",
    "Can you refinance into an AIO loan in Texas",
    "Why the AIO loan is one of the most powerful loan products in Texas",
    "How couples and dual-income households use AIO loans to build wealth",

    # New construction and builders (Houston metro focus)
    "Buying new construction in the Houston metro — buyer's complete guide",
    "Best new home builders in the Houston metro 2026",
    "Why you should use your own mortgage lender on a Houston new build",
    "Two-Time Close construction loans in Texas — what to know",
    "New construction vs resale in the Houston area — the real tradeoffs",
    "What to negotiate when buying new construction in the Houston metro",
    "Why builder incentives can hide bad mortgage terms",
    "MUD bonds and your new construction Texas home",
    "What inspectors look for in new construction homes in Houston",
    "Buyer's checklist for new construction in the Houston metro",

    # Spec construction financing for builders
    "How builders use spec financing to scale in the Houston market",
    "What is a Spec Lock and how it protects Houston builders",
    "Why Texas builders need a strong mortgage lender partner",
    "Construction loan draws explained for Houston builders",
    "How builders qualify buyers for FHA spec homes in Texas",
    "Why VA buyers love new construction in the Houston metro market",
    "How HomeStyle renovation loans work in the Houston area",
    "Why builders should send their buyers to Tyler Henschel",
    "Spec vs custom construction loans in Texas — which is right for your project",
    "How Houston builders use spec financing to scale inventory",
    "What builders need to know about construction capital in the Houston market",
    "How to finance land development and vertical construction in Texas",

    # Realtor and referral SEO
    "How realtors and lenders work together in Houston TX",
    "What makes a great mortgage lender for Houston realtors",
    "How Houston realtors can build a stronger referral business",
    "Why buyers without realtors still need a great mortgage lender in Texas",
    "What every Houston realtor should know about mortgage pre-approval",
    "How a strong lender relationship helps Houston realtors close more deals",
    "Why having the right lender matters more than the interest rate in Texas",
    "What makes a great loan officer in Houston TX",
]

# Allow DAY_OFFSET env to generate a future-dated draft (used for batch
# weekly generation: 0 = today, 1 = tomorrow, ..., 6 = next week).
day_offset = int(os.environ.get("DAY_OFFSET", "0"))
target_date = datetime.datetime.now() + datetime.timedelta(days=day_offset)
day_of_year = target_date.timetuple().tm_yday
topic = TOPICS[day_of_year % len(TOPICS)]
date_str = target_date.strftime("%B %d, %Y")
date_iso = target_date.strftime("%Y-%m-%d")

slug = topic.lower()
slug = re.sub(r'[^a-z0-9\s]', '', slug)
slug = re.sub(r'\s+', '-', slug.strip())
slug = slug[:60].rstrip('-')

draft_dir = f"_drafts/{date_iso}_{slug}"
print(f"Topic: {topic}")
print(f"Draft folder: {draft_dir}")

# ---------------------------------------------------------------------------
# 1. FAQ generation
# ---------------------------------------------------------------------------
faq_message = client.messages.create(
    model="claude-haiku-4-5-20251001",
    max_tokens=600,
    messages=[{
        "role": "user",
        "content": f"""Generate 3 FAQ questions and answers about: {topic}

Format as JSON array only, no other text:
[
  {{"q": "Question one?", "a": "Answer one."}},
  {{"q": "Question two?", "a": "Answer two."}},
  {{"q": "Question three?", "a": "Answer three."}}
]

Keep answers 2-3 sentences. Focus on what Houston-area homebuyers actually search for. Rotate across Houston metro communities (Tomball, Cypress, Sugar Land, The Woodlands, Pearland, Spring, Magnolia, Conroe, Katy) — do not lean exclusively on Katy."""
    }]
)

try:
    faq_data = json.loads(faq_message.content[0].text.strip())
except Exception:
    faq_data = []

faq_schema = ""
if faq_data:
    faq_items = ",\n".join([
        f'{{"@type":"Question","name":{json.dumps(f["q"])},"acceptedAnswer":{{"@type":"Answer","text":{json.dumps(f["a"])}}}}}'
        for f in faq_data
    ])
    faq_schema = f"""<script type="application/ld+json">
{{"@context":"https://schema.org","@type":"FAQPage","mainEntity":[{faq_items}]}}
</script>"""

faq_html = ""
if faq_data:
    faq_items_html = "\n".join([
        f'<div class="faq-item"><h3>{f["q"]}</h3><p>{f["a"]}</p></div>'
        for f in faq_data
    ])
    faq_html = f'<div class="faq-block"><h2>Frequently Asked Questions</h2>{faq_items_html}</div>'

# ---------------------------------------------------------------------------
# 2. Blog body generation
# ---------------------------------------------------------------------------
message = client.messages.create(
    model="claude-haiku-4-5-20251001",
    max_tokens=2000,
    messages=[{
        "role": "user",
        "content": f"""Write an SEO-optimized blog post for tylerhloans.com about: {topic}

Tyler Henschel is a Senior Loan Officer at CMG Home Loans serving the Greater Houston area (NMLS 2034073).
Phone: (979) 255-6219
Email: thenschel@cmghomeloans.com
Calendly: https://calendly.com/thenschel-cmghomeloans/30min

Requirements:
- 800 to 1000 words
- Conversational but professional tone
- Target keywords naturally
- Focus on the Houston MSA broadly. Tyler's office is in Katy but he serves the entire Houston metro — Tomball, Cypress, Sugar Land, The Woodlands, Pearland, Spring, Magnolia, Conroe, Friendswood, Memorial, Heights, Bellaire, and beyond. Do NOT lean on Katy as the default. Rotate which specific Houston neighborhood you mention based on the topic.
- Strong CTA at the end to book a call with Tyler
- Format as clean HTML using h2, h3, p, ul, li tags only
- Start with a meta description comment like: <!-- META: description here -->
- Do NOT include DOCTYPE html head or body tags
- Mention Tyler naturally 2 to 3 times as the expert

COMPLIANCE RULES (CMG Financial Marketing Policy — strictly follow all of these):
- NEVER quote specific interest rates or APR figures. General language like "competitive rates" is fine.
- NEVER imply or state that any reader is pre-approved, guaranteed, or likely to qualify for any loan.
- NEVER use the word "fixed" in a way that could mislead readers about rate variability.
- NEVER make claims about eliminating, reducing, or restructuring debt.
- NEVER misrepresent any loan product as government-backed or government-affiliated unless it explicitly is (e.g., VA, FHA, USDA).
- NEVER state or imply that no fees or costs are associated with any loan product.
- NEVER make any claim about specific payment amounts without full context.
- Do NOT display or reference any specific current mortgage rates — rates change daily and must come from an approved source.
- All product descriptions must be accurate and not misleading. Do not overstate benefits.
- Use educational, informational language throughout. The post is informational, not a binding offer of credit.
- NEVER use unqualified superlatives or unprovable claims such as "best rates," "lowest costs," "cheapest," or "free" (unless a specific product genuinely has no fee and that is documented).
- Keep content general and educational in nature. Do not frame the post as a specific product advertisement with rate or payment terms.
- Use inclusive, welcoming language. Never use phrasing that could discourage any person from applying based on race, color, national origin, religion, sex, familial status, or disability (Fair Housing Act / ECOA compliance).
- Do not make any statement that is inaccurate, inconsistent, or could be considered deceptive or misleading under UDAAP standards.

PRODUCT ACCURACY — CMG GUIDELINES YOU MUST RESPECT:
- The All-In-One (AIO) loan is secured against a PRIMARY RESIDENCE or SECOND HOME only. Investment / rental properties are NOT eligible for AIO. Do not say or imply otherwise.
- The AIO is a 30-year revolving line of credit. It does NOT have a separate "draw period followed by repayment period" the way a HELOC does. Interest accrues on the daily net balance.
- CMG's conventional construction product uses a Two-Time Close structure (not single close). Minimum FICO is 720 for ground-up construction. The 680 minimum applies to renovation products only (HomeStyle, etc.), not ground-up construction.
- For DSCR loans: the formula is Gross Rent / PITIA (monthly), NOT NOI / annual debt service. Gross rent comes from Form 1007/1025 or the lease — not NOI. Long-term DSCR minimums commonly range ~0.75–1.15. The 1.25 figure is specific to short-term-rental (STR) programs.
- Texas A6 home equity rules: cash-out is capped at 80% LTV (vs 100% purchase, 95% rate/term). 12-day cooling-off period before closing. 1-year seasoning since the most recent prior equity loan. Disclosure notices in English and Spanish.
- For DTI: Conventional purchase commonly works around 43–50% DTI depending on AUS. Non-QM (Edge, Sharp, Advantage) can accommodate up to 55%. CMG's conventional construction product caps back-end DTI at 45% and front-end at 38%. FHA TOTAL Scorecard routinely approves above 43% DTI.
- PMI (conventional, removable at 20% equity) is DIFFERENT from FHA MIP (lifetime for most post-2013 FHA loans <10% down; only path to remove is refinance out of FHA). Never conflate the two.
- Spec construction financing targets experienced builders — typically 5 to 50 homes per year with a proven track record.
- The Spec Lock (Build and Lock) program is available for Conventional, FHA, VA, and USDA loans only. It is NOT available for Non-Agency products.
- CRITICAL: The All-In-One (AIO) loan is NOT available for spec construction financing in Texas. Never suggest or imply that Texas builders can use the AIO loan for spec construction purposes."""
    }]
)

content = message.content[0].text

meta_match = re.search(r'<!-- META: (.*?) -->', content)
meta_desc = meta_match.group(1) if meta_match else f"Tyler Henschel, Senior Loan Officer serving the Greater Houston area, explains everything about {topic.lower()}."

# ---------------------------------------------------------------------------
# 3. 60-second video script generation (for Tyler to record)
# ---------------------------------------------------------------------------
script_message = client.messages.create(
    model="claude-haiku-4-5-20251001",
    max_tokens=400,
    messages=[{
        "role": "user",
        "content": f"""Write a 60-second video script for Tyler Henschel (Senior Loan Officer, CMG Home Loans, NMLS #2034073) to read on camera. The video accompanies a blog post titled: "{topic}"

Requirements:
- 150 to 160 words exactly. Spoken pace is roughly 150 words per minute.
- First person — Tyler talking directly to camera. Use "I", "me", "my".
- HOOK in the first 7 seconds — get attention immediately. A surprising fact, common mistake, or pointed question.
- Hit 2-3 main takeaways from the blog topic. Don't try to cover everything.
- Soft CTA at the end: book a call at calendly.com/thenschel-cmghomeloans/30min OR visit tylerhloans.com
- Conversational and natural — short sentences, contractions, real rhythm. Avoid corporate filler ("In today's video we will discuss"). Talk like a real person.
- Tyler serves the entire Houston metro — Tomball, Cypress, Sugar Land, The Woodlands, Pearland, Spring, Magnolia, Conroe, Katy. Rotate which area you name based on topic relevance.
- NO specific interest rates, APR figures, or "best/lowest/cheapest" superlatives. Compliance: do not imply pre-approval or guarantee.
- AIO is for primary or second home only (not investment properties).
- Do not include camera directions, stage notes, or "[pause]" markers. Just the words Tyler says.

Output ONLY the script body. No intro line, no labels, no JSON. Just the words."""
    }]
)
video_script = script_message.content[0].text.strip()

# ---------------------------------------------------------------------------
# 4. Pre-render the blog page HTML (publisher will move + add YouTube embed)
# ---------------------------------------------------------------------------
full_page = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{topic} | Tyler Henschel CMG Home Loans Houston TX</title>
<meta name="description" content="{meta_desc}">
<link rel="canonical" href="https://tylerhloans.com/blog/{slug}">
<meta property="og:title" content="{topic} | Tyler Henschel">
<meta property="og:description" content="{meta_desc}">
<meta property="og:url" content="https://tylerhloans.com/blog/{slug}">
<meta property="og:type" content="article">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,400;1,400&family=DM+Sans:wght@300;400;500&display=swap" rel="stylesheet">
<script type="application/ld+json">
{{"@context":"https://schema.org","@type":"Article","headline":{json.dumps(topic)},"description":{json.dumps(meta_desc)},"datePublished":"{date_str}","author":{{"@type":"Person","name":"Tyler Henschel","url":"https://tylerhloans.com","jobTitle":"Senior Loan Officer","worksFor":{{"@type":"Organization","name":"CMG Home Loans"}}}},"publisher":{{"@type":"Organization","name":"CMG Home Loans","url":"https://tylerhloans.com"}}}}
</script>
{faq_schema}
<style>
:root{{--navy:#0a1628;--gold:#c9a84c;--gold-light:#e8c97a;--cream:#f7f4ef;--text:#1a1a2e;--muted:#5a6070;--white:#ffffff}}
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'DM Sans',sans-serif;background:var(--cream);color:var(--text)}}
nav{{background:var(--navy);padding:1rem 2rem;display:flex;justify-content:space-between;align-items:center}}
.nav-name{{font-family:'Playfair Display',serif;color:var(--gold-light);font-size:1rem;text-decoration:none}}
.nav-cta{{background:var(--gold);color:var(--navy);padding:.5rem 1.2rem;border-radius:4px;text-decoration:none;font-size:.8rem;font-weight:500}}
.hero{{background:var(--navy);padding:4rem 2rem;text-align:center}}
.hero h1{{font-family:'Playfair Display',serif;color:var(--white);font-size:clamp(1.6rem,3vw,2.6rem);font-weight:400;max-width:800px;margin:0 auto 1rem;line-height:1.2}}
.hero-meta{{color:rgba(255,255,255,.45);font-size:.85rem}}
.article{{max-width:780px;margin:3rem auto;padding:0 2rem 4rem}}
.video-embed{{position:relative;padding-bottom:56.25%;height:0;overflow:hidden;border-radius:12px;margin-bottom:2.5rem;background:#000}}
.video-embed iframe{{position:absolute;top:0;left:0;width:100%;height:100%;border:0}}
.article h2{{font-family:'Playfair Display',serif;font-size:1.5rem;font-weight:400;color:var(--navy);margin:2rem 0 .8rem}}
.article h3{{font-size:1.05rem;font-weight:500;color:var(--navy);margin:1.5rem 0 .5rem}}
.article p{{font-size:.97rem;font-weight:300;line-height:1.8;color:var(--muted);margin-bottom:1.1rem}}
.article ul,.article ol{{padding-left:1.5rem;margin-bottom:1.1rem}}
.article li{{font-size:.97rem;font-weight:300;line-height:1.8;color:var(--muted);margin-bottom:.3rem}}
.article strong{{font-weight:500;color:var(--text)}}
.cta-box{{background:var(--navy);border-radius:16px;padding:2.5rem;text-align:center;margin-top:3rem}}
.cta-box h3{{font-family:'Playfair Display',serif;color:var(--white);font-size:1.4rem;font-weight:400;margin-bottom:.8rem}}
.cta-box p{{color:rgba(255,255,255,.55);margin-bottom:1.5rem;font-weight:300}}
.cta-btn{{background:var(--gold);color:var(--navy);padding:.9rem 2rem;border-radius:4px;text-decoration:none;font-weight:500;display:inline-block}}
.back{{display:block;text-align:center;margin-top:2rem;color:var(--muted);text-decoration:none;font-size:.85rem}}
.author-bio{{background:var(--white);border:1px solid rgba(10,22,40,0.08);border-radius:16px;padding:2rem;margin-top:3rem;display:grid;grid-template-columns:140px 1fr;gap:1.8rem;align-items:center}}
.author-bio-photo{{width:140px;height:140px;border-radius:50%;overflow:hidden;border:3px solid var(--gold);flex-shrink:0}}
.author-bio-photo img{{width:100%;height:100%;object-fit:cover;display:block}}
.author-bio-label{{font-size:0.65rem;font-weight:500;letter-spacing:0.14em;text-transform:uppercase;color:var(--gold);margin-bottom:0.4rem}}
.author-bio h3{{font-family:'Playfair Display',serif;font-size:1.3rem;font-weight:400;color:var(--navy);margin-bottom:0.2rem}}
.author-bio-title{{font-size:0.82rem;color:var(--muted);margin-bottom:0.8rem;font-weight:300}}
.author-bio-desc{{font-size:0.85rem;font-weight:300;color:var(--muted);line-height:1.65;margin-bottom:1rem}}
.author-bio-contact{{display:flex;flex-wrap:wrap;gap:0.8rem 1.5rem;font-size:0.78rem}}
.author-bio-contact a{{color:var(--navy);text-decoration:none;font-weight:500;display:inline-flex;align-items:center;gap:0.4rem}}
.author-bio-contact a:hover{{color:var(--gold)}}
.author-bio-nmls{{font-size:0.7rem;color:var(--muted);margin-top:0.8rem;font-weight:300;letter-spacing:0.04em}}
@media(max-width:600px){{.author-bio{{grid-template-columns:1fr;text-align:center;padding:1.8rem}}.author-bio-photo{{margin:0 auto}}.author-bio-contact{{justify-content:center}}}}
footer{{background:#060e1b;padding:2rem;text-align:center}}
footer p{{font-size:.7rem;color:rgba(255,255,255,.22);line-height:1.6;margin-bottom:.2rem}}
</style>
</head>
<body>
<nav style="position:fixed;top:0;left:0;right:0;z-index:100;display:flex;align-items:center;justify-content:space-between;padding:1rem 2.5rem;background:rgba(10,22,40,0.97);backdrop-filter:blur(12px);border-bottom:1px solid rgba(201,168,76,0.15);flex-wrap:wrap;gap:0.8rem">
<a style="text-decoration:none" href="/"><div style="font-family:'Playfair Display',serif;font-size:1rem;color:#e8c97a">Tyler Henschel</div><div style="font-size:0.65rem;color:rgba(255,255,255,0.35);letter-spacing:0.08em">CMG Home Loans · NMLS #2034073</div></a>
<div style="display:flex;gap:1.3rem;align-items:center;flex-wrap:wrap">
<a style="color:rgba(255,255,255,0.65);text-decoration:none;font-size:0.82rem" href="/">Home</a>
<a style="color:rgba(255,255,255,0.65);text-decoration:none;font-size:0.82rem" href="/products/">Products</a>
<a style="color:rgba(255,255,255,0.65);text-decoration:none;font-size:0.82rem" href="/investors/">Investors</a>
<a style="color:rgba(255,255,255,0.65);text-decoration:none;font-size:0.82rem" href="/builders/">Builders</a>
<a style="color:rgba(255,255,255,0.65);text-decoration:none;font-size:0.82rem" href="/refinance/">Refinance</a>
<a style="color:#e8c97a;text-decoration:none;font-size:0.82rem" href="/blog/">Blog</a>
<a style="background:transparent;color:#c9a84c;font-size:0.78rem;font-weight:500;padding:0.45rem 1rem;border-radius:4px;text-decoration:none;border:1px solid rgba(201,168,76,0.4)" href="https://www.cmghomeloans.com/mysite/tyler-henschel" target="_blank">Apply Now</a>
<a style="background:#c9a84c;color:#0a1628;font-size:0.78rem;font-weight:500;padding:0.5rem 1.2rem;border-radius:4px;text-decoration:none" href="https://calendly.com/thenschel-cmghomeloans/30min" target="_blank">Book a Call</a>
</div>
</nav>
<div class="hero">
<h1>{topic}</h1>
<p class="hero-meta">By Tyler Henschel &middot; {date_str} &middot; Houston, TX</p>
</div>
<article class="article">
<!-- VIDEO_EMBED_PLACEHOLDER -->
{content}
{faq_html}
<style>.faq-block{{background:#f7f4ef;border-radius:12px;padding:2rem;margin:2rem 0}}.faq-block h2{{font-family:'Playfair Display',serif;font-size:1.4rem;font-weight:400;color:#0a1628;margin-bottom:1.2rem}}.faq-item{{margin-bottom:1.2rem;padding-bottom:1.2rem;border-bottom:1px solid rgba(10,22,40,0.08)}}.faq-item:last-child{{margin-bottom:0;padding-bottom:0;border-bottom:none}}.faq-item h3{{font-size:0.95rem;font-weight:500;color:#0a1628;margin-bottom:0.4rem}}.faq-item p{{font-size:0.88rem;font-weight:300;color:#5a6070;line-height:1.7}}</style>
<div class="author-bio">
<div class="author-bio-photo"><img src="/images/tyler-henschel.jpg" alt="Tyler Henschel Senior Loan Officer CMG Home Loans Houston Katy TX"></div>
<div>
<div class="author-bio-label">About the Author</div>
<h3>Tyler Henschel</h3>
<div class="author-bio-title">Senior Loan Officer &middot; CMG Home Loans</div>
<p class="author-bio-desc">Tyler serves homebuyers, investors, and builders across the greater Houston area &mdash; Tomball, Cypress, Sugar Land, The Woodlands, Pearland, Spring, Magnolia, Conroe, Katy, and beyond. From first-time buyers to seasoned investors scaling DSCR portfolios, his goal is straightforward: get you the right loan, the first time.</p>
<div class="author-bio-contact">
<a href="tel:9792556219">📞 (979) 255-6219</a>
<a href="mailto:thenschel@cmghomeloans.com">✉️ thenschel@cmghomeloans.com</a>
<a href="https://calendly.com/thenschel-cmghomeloans/30min" target="_blank">📅 Book a Call</a>
<a href="https://www.cmghomeloans.com/mysite/tyler-henschel" target="_blank">📝 Apply Now</a>
</div>
<div class="author-bio-nmls">NMLS #2034073 &middot; 24285 Katy Freeway Suite 400-G, Katy TX 77494</div>
</div>
</div>
<div class="cta-box">
<h3>Ready to explore your options?</h3>
<p>Tyler Henschel specializes in helping Houston-area borrowers find the right mortgage strategy. Book a free 20-minute call.</p>
<a class="cta-btn" href="https://calendly.com/thenschel-cmghomeloans/30min" target="_blank">Book a Free Strategy Call</a>
</div>
<a class="back" href="/blog">&larr; Back to all articles</a>
</article>
<footer>
<p>Tyler Henschel &middot; Senior Loan Officer &middot; CMG Home Loans &middot; NMLS #2034073</p>
<p>24285 Katy Freeway Suite 400-G, Katy, TX 77494 &middot; (979) 255-6219 &middot; thenschel@cmghomeloans.com</p>
<p>Offer of credit subject to credit approval. CMG Mortgage, Inc., dba CMG Financial NMLS #1820 is an equal opportunity lender.</p>
<p>&copy; 2025 CMG Home Loans. Not a commitment to lend. Equal Housing Lender.</p>
</footer>
</body>
</html>"""

os.makedirs(draft_dir, exist_ok=True)
with open(f"{draft_dir}/index.html", "w") as f:
    f.write(full_page)
print(f"Wrote {draft_dir}/index.html")

# Save metadata (publisher reads this when promoting to live)
meta_payload = {
    "topic": topic,
    "slug": slug,
    "date": date_str,
    "date_iso": date_iso,
    "meta_desc": meta_desc,
}
with open(f"{draft_dir}/meta.json", "w") as f:
    json.dump(meta_payload, f, indent=2)

# Save raw body content (publisher may use it for regenerated HTML if needed)
with open(f"{draft_dir}/body.html", "w") as f:
    f.write(content)

# ---------------------------------------------------------------------------
# 5. Build the review.docx file Tyler proofreads
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        """Convert simple blog HTML into ordered paragraphs/headings for docx."""
        def __init__(self):
            super().__init__()
            self.blocks = []        # list of (style, text)
            self.buf = []
            self.current_style = "Normal"
            self.in_list = False
            self.li_buf = []

        def _flush(self):
            text = "".join(self.buf).strip()
            if text:
                self.blocks.append((self.current_style, text))
            self.buf = []

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag in ("h1", "h2"):
                self._flush()
                self.current_style = "Heading 2"
            elif tag == "h3":
                self._flush()
                self.current_style = "Heading 3"
            elif tag == "h4":
                self._flush()
                self.current_style = "Heading 4"
            elif tag == "p":
                self._flush()
                self.current_style = "Normal"
            elif tag in ("ul", "ol"):
                self.in_list = True
            elif tag == "li":
                self._flush()
                self.current_style = "List Bullet"

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in ("h1", "h2", "h3", "h4", "p", "li"):
                self._flush()
                self.current_style = "Normal"
            elif tag in ("ul", "ol"):
                self.in_list = False

        def handle_data(self, data):
            self.buf.append(data)

    doc = Document()

    # ---- Header section ----
    title_p = doc.add_heading(topic, level=1)

    meta_lines = [
        f"Publish date: {date_str}",
        f"Slug: {slug}",
        "Status: APPROVED   |   Mark DELETE or EDIT here: ___________________",
        "Paste YouTube URL here (after recording + Submagic): ___________________",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer

    # ---- Video script section ----
    doc.add_heading("🎬 60-Second Video Script", level=2)
    p = doc.add_paragraph(video_script)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # spacer

    # ---- Blog body section ----
    doc.add_heading("📝 Blog Post", level=2)
    parser = _TextExtractor()
    parser.feed(content)
    parser._flush()
    for style, text in parser.blocks:
        if style == "Normal":
            doc.add_paragraph(text)
        else:
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                doc.add_paragraph(text)

    # ---- FAQ section in docx ----
    if faq_data:
        doc.add_paragraph()
        doc.add_heading("FAQ", level=2)
        for f in faq_data:
            p = doc.add_paragraph()
            p.add_run(f["q"]).bold = True
            doc.add_paragraph(f["a"])

    doc.save(f"{draft_dir}/review.docx")
    print(f"Wrote {draft_dir}/review.docx")

except ImportError:
    print("python-docx not installed; skipping review.docx")

# ---------------------------------------------------------------------------
# 6. Commit + push in CI (with retry to handle concurrent runs)
# ---------------------------------------------------------------------------
if os.environ.get("CI"):
    import time
    subprocess.run(["git", "config", "user.email", "blog-agent@tylerhloans.com"], check=True)
    subprocess.run(["git", "config", "user.name", "Blog Agent"], check=True)
    subprocess.run(["git", "add", draft_dir], check=True)
    subprocess.run(["git", "commit", "-m", f"draft: {topic[:72]}"], check=True)

    # Retry push up to 5 times with pull-rebase between attempts (handles
    # concurrent workflow runs from batch weekly generation).
    for attempt in range(5):
        push = subprocess.run(["git", "push"], capture_output=True, text=True)
        if push.returncode == 0:
            print(f"Pushed draft: {draft_dir}")
            break
        print(f"Push attempt {attempt+1} failed: {push.stderr.strip()[:200]}")
        time.sleep(2 + attempt * 2)
        pull = subprocess.run(["git", "pull", "--rebase", "origin", "main"], capture_output=True, text=True)
        if pull.returncode != 0:
            print(f"Rebase failed: {pull.stderr.strip()[:200]}")
            # If rebase fails, try without rebase
            subprocess.run(["git", "rebase", "--abort"], capture_output=True)
            subprocess.run(["git", "pull", "origin", "main"], capture_output=True)
    else:
        raise RuntimeError("Could not push draft after 5 attempts")
